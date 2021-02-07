import docx
import time
import re

#################################################################################
########################## ENGLISH HOMEWORK SOLVER ##############################
#################################################################################

def solve(text):
    if "if" or "If" in text:
        try:
            keyword = "if"
            before, keyword, after = text.partition(keyword)
            IfFirstCondition = re.compile(r'\w+')
            query = after.split(" ")
            query.pop(1)
            after = " ".join(query)
            print(after)
            found = IfFirstCondition.search(after)
            phrase = found.group()
            verbEnd = phrase[-1:]
            print(verbEnd, phrase)
            if "had" in phrase:
                print("If third conditional, Answer: would have")
            else:
                if verbEnd == "d":
                    print("If second conditional, Answer: would")
                elif verbEnd == "s":
                    print("If first conditional, Answer: will")
                else:
                    print("If first conditional, Answer: will")
        except:
            keyword = "If"
            before, keyword, after = text.partition(keyword)
            IfFirstCondition = re.compile(r'\w+')
            query = after.split(" ")
            query.pop(1)
            after = " ".join(query)
            print(after)
            found = IfFirstCondition.search(after)
            phrase = found.group()
            verbEnd = phrase[-1:]
            print(verbEnd, phrase)
            if "had" in phrase:
                print("If third conditional, Answer: would have")
            else:
                if verbEnd == "d":
                    print("If second conditional, Answer: would")
                elif verbEnd == "s":
                    print("If first conditional, Answer: will")
                else:
                    print("If first conditional, Answer: will")


doc = docx.Document("solve.docx")
print(len(doc.paragraphs))
for i in range(0, len(doc.paragraphs),2):
               solve(doc.paragraphs[i].text)
